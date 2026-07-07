from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

UNSUPPORTED_FILE_TYPE_ERROR = "UNSUPPORTED_FILE_TYPE"
PDF_EXTRACTION_ERROR = "PDF_TEXT_EXTRACTION_FAILED"
DOCX_EXTRACTION_ERROR = "DOCX_TEXT_EXTRACTION_FAILED"
TEXT_EXTRACTION_ERROR = "TEXT_EXTRACTION_FAILED"


def extract_text(filename: str, content: bytes) -> tuple[str, str | None]:
    extension = Path(filename).suffix.lower()
    try:
        if extension == ".pdf":
            return extract_pdf_text(content), None
        if extension == ".docx":
            return extract_docx_text(content), None
        return "", UNSUPPORTED_FILE_TYPE_ERROR
    except Exception:
        return "", safe_extraction_error(extension)


def safe_extraction_error(extension: str) -> str:
    if extension == ".pdf":
        return PDF_EXTRACTION_ERROR
    if extension == ".docx":
        return DOCX_EXTRACTION_ERROR
    return TEXT_EXTRACTION_ERROR


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    page_texts = []
    for page in reader.pages:
        page_texts.append(page.extract_text() or "")
    return "\n".join(page_texts).strip()


def extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts).strip()
